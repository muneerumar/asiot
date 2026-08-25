"""Build the polished Word technical guide from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ASIoT_Simulation_Code_Documentation.md"
OUTPUT = ROOT / "docs" / "ASIoT_Simulation_Code_Documentation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
GOLD = "A36B00"
TABLE_FILL = "E8EEF5"
CODE_FILL = "F4F6F9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
PAGE_BREAK_HEADINGS = {
    "3. Repository map",
    "4. Configuration and workload",
    "16. Known limitations and manuscript actions",
}


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    _configure_document(doc)
    bullet_id, number_id = _configure_numbering(doc)
    _add_cover(doc)
    _add_contents(doc, text)
    _render_markdown(doc, text, bullet_id, number_id)
    _set_core_properties(doc)
    doc.save(OUTPUT)
    _audit(OUTPUT)
    print(OUTPUT)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    _set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        _set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("A-SIoT SIMULATION  |  TECHNICAL GUIDE")
    _format_run(run, "Calibri", 8.5, MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    label = p.add_run("PAGE ")
    _format_run(label, "Calibri", 8.5, MUTED)
    _append_field(p, "PAGE")


def _add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(105)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("TECHNICAL AND REPRODUCIBILITY GUIDE")
    _format_run(run, "Calibri", 10, GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("A-SIoT Simulation")
    _format_run(run, "Calibri", 30, NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("Code Architecture, Evaluated Protocols, and Validated Results")
    _format_run(run, "Calibri", 15, DARK_BLUE)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(78)
    run = tagline.add_run("A compact reference for reproduction, review, and manuscript correction")
    _format_run(run, "Calibri", 10.5, MUTED, italic=True)

    for value, bold in (
        ("Revision  |  21 August 2026", True),
        ("Status  |  Research artifact validated; manuscript integration pending", False),
        ("Evidence  |  supplementary_results/", False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(value)
        _format_run(run, "Calibri", 10, NAVY if bold else MUTED, bold=bold)
    doc.add_page_break()


def _add_contents(doc: Document, markdown: str) -> None:
    doc.add_heading("Contents", level=1)
    intro = doc.add_paragraph(
        "This guide is organized around the implemented runtime, the evaluated studies, "
        "and the remaining manuscript-only actions."
    )
    intro.paragraph_format.space_after = Pt(10)
    headings = [
        match.group(1).strip()
        for match in re.finditer(r"^##\s+(.+)$", markdown, flags=re.MULTILINE)
    ]
    for index, heading in enumerate(headings, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{index:02d}   {heading}")
        _format_run(run, "Calibri", 10.5, DARK_BLUE, bold=index in {1, 2, 13, 15})
    doc.add_page_break()


def _render_markdown(doc: Document, markdown: str, bullet_id: int, number_id: int) -> None:
    lines = markdown.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## "))
    index = start
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            _add_code_block(doc, "\n".join(code_lines))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [stripped]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_table(doc, [_parse_table_row(row) for row in table_lines])
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            index += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading in PAGE_BREAK_HEADINGS:
                doc.add_page_break()
            doc.add_heading(heading, level=1)
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            text = (bullet or numbered).group(1)
            continuation, index = _consume_continuation(lines, index + 1)
            _add_list_item(
                doc,
                " ".join([text] + continuation),
                bullet_id if bullet else number_id,
            )
            continue
        if stripped == "---":
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or _starts_block(lines, index):
                break
            paragraph_lines.append(candidate)
            index += 1
        _add_rich_paragraph(doc, " ".join(paragraph_lines))


def _add_rich_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    _add_inline_runs(p, text)
    return p


def _add_list_item(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    _add_inline_runs(p, text)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    _set_paragraph_shading(p, CODE_FILL)
    run = p.add_run(code or " ")
    _format_run(run, "Consolas", 8.5, "303642")


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    widths = _column_widths(columns)
    table = doc.add_table(rows=len(normalized), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_geometry(table, widths)
    for row_index, row in enumerate(normalized):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_width(cell, widths[column_index])
            _set_cell_margins(cell, 80, 80, 120, 120)
            if row_index == 0:
                _set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            _add_inline_runs(p, value, base_size=8.5, force_bold=row_index == 0)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _column_widths(columns: int) -> list[int]:
    patterns = {
        1: [9360],
        2: [2700, 6660],
        3: [2100, 3000, 4260],
        4: [1500, 2160, 2160, 3540],
        5: [1500, 1900, 1900, 1900, 2160],
    }
    if columns in patterns:
        return patterns[columns]
    base = TABLE_WIDTH_DXA // columns
    values = [base] * columns
    values[-1] += TABLE_WIDTH_DXA - sum(values)
    return values


def _configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    next_abstract = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1
    bullet_id = _append_numbering(numbering, next_abstract, next_num, "bullet", "•")
    number_id = _append_numbering(numbering, next_abstract + 1, next_num + 1, "decimal", "%1.")
    return bullet_id, number_id


def _append_numbering(numbering, abstract_id: int, num_id: int, fmt: str, text: str) -> int:
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, indent, spacing))
    level.extend((start, num_fmt, lvl_text, lvl_jc, p_pr))
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _set_table_geometry(table, widths: list[int]) -> None:
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def _set_cell_width(cell, width: int) -> None:
    cell.width = Inches(width / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_inline_runs(p, text: str, base_size: float = 11, force_bold: bool = False) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = p.add_run(text[position:match.start()])
            _format_run(run, "Calibri", base_size, "000000", bold=force_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2])
            _format_run(run, "Calibri", base_size, "000000", bold=True)
        else:
            run = p.add_run(token[1:-1])
            _format_run(run, "Consolas", max(8, base_size - 1), DARK_BLUE, bold=force_bold)
        position = match.end()
    if position < len(text):
        run = p.add_run(text[position:])
        _format_run(run, "Calibri", base_size, "000000", bold=force_bold)


def _format_run(run, font: str, size: float, color: str, *, bold=False, italic=False) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _set_style_font(style, font: str, size: float, color: str, bold=False) -> None:
    style.font.name = font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _append_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(value)
    run.append(end)
    paragraph._p.append(run)


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _starts_block(lines: list[str], index: int) -> bool:
    stripped = lines[index].strip()
    return bool(
        stripped.startswith(("## ", "### ", "```", "|", "- "))
        or re.match(r"^\d+\.\s+", stripped)
        or stripped == "---"
    )


def _consume_continuation(lines: list[str], index: int) -> tuple[list[str], int]:
    values: list[str] = []
    while index < len(lines):
        raw = lines[index]
        if not raw.strip() or _starts_block(lines, index):
            break
        values.append(raw.strip())
        index += 1
    return values, index


def _set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "A-SIoT Simulation: Code and Reproducibility Guide"
    props.subject = "Architecture, protocols, validated results, and limitations"
    props.author = "A-SIoT Simulation Project"
    props.keywords = "A-SIoT, simulation, trust, privacy, attacks, MARL, reproducibility"


def _audit(path: Path) -> None:
    check = Document(path)
    section = check.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    assert all(
        round(value.inches, 3) == 1.0
        for value in (
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        )
    )
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    assert check.styles["Normal"].font.name == "Calibri"
    for table in check.tables:
        width = table._tbl.tblPr.find(qn("w:tblW"))
        indent = table._tbl.tblPr.find(qn("w:tblInd"))
        assert width.get(qn("w:w")) == str(TABLE_WIDTH_DXA)
        assert width.get(qn("w:type")) == "dxa"
        assert indent.get(qn("w:w")) == str(TABLE_INDENT_DXA)
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == TABLE_WIDTH_DXA
        for row in table.rows:
            assert [
                int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w")))
                for cell in row.cells
            ] == grid_widths


if __name__ == "__main__":
    main()
